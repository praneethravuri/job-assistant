import json
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


class ResumeJsonToDocx:
    """
    Builds a .docx resume from given JSON data.
    """

    def __init__(self, resume_data: dict):
        self.data = resume_data

    @staticmethod
    def _add_horizontal_line(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _make_two_column_paragraph(doc, left_text, right_text, bold_left=False, italic_left=False):
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1
        para.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
        run_left = para.add_run(left_text)
        if bold_left:
            run_left.bold = True
        if italic_left:
            run_left.italic = True
        para.add_run(f'\t{right_text}')
        return para

    @staticmethod
    def _add_bulleted_item(doc, text):
        para = doc.add_paragraph(style='List Bullet')
        fmt = para.paragraph_format
        fmt.left_indent = Inches(0.25)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1
        para.add_run(text)
        return para

    def generate(self, output_path: str = "Resume.docx"):
        doc = Document()
        sec = doc.sections[0]
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)
        sec.top_margin = Inches(0.25)
        sec.bottom_margin = Inches(0.25)
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1

        # Header
        hdr = self.data.get('header', {})
        name_p = doc.add_paragraph()
        name_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = name_p.add_run(hdr.get('name', ''))
        run.bold = True
        run.font.size = Pt(16)
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_info = ' | '.join(hdr.get('contact', []))
        contact_p.add_run(contact_info).font.size = Pt(10)

        # Work Experience
        we_title = doc.add_paragraph()
        we_title.add_run('WORK EXPERIENCE').bold = True
        we_title.runs[0].font.size = Pt(12)
        self._add_horizontal_line(we_title)
        for exp in self.data.get('work_experience', []):
            p1 = self._make_two_column_paragraph(doc, exp.get('company', ''), exp.get('location', ''), bold_left=True)
            p1.space_before = Pt(4)
            p2 = self._make_two_column_paragraph(doc, exp.get('position', ''), f"{exp.get('start_date', '')} – {exp.get('end_date', '')}", bold_left=True)
            p2.space_after = Pt(4)
            for bullet in exp.get('bullets', []):
                self._add_bulleted_item(doc, bullet)

        # Education
        edu_title = doc.add_paragraph()
        edu_title.add_run('EDUCATION').bold = True
        edu_title.runs[0].font.size = Pt(12)
        self._add_horizontal_line(edu_title)
        for edu in self.data.get('education', []):
            self._make_two_column_paragraph(doc, edu.get('institution', ''), edu.get('location', ''), bold_left=True)
            self._make_two_column_paragraph(doc, edu.get('degree', ''), f"{edu.get('start_date', '')} – {edu.get('end_date', '')}")

        # Skills
        skills_title = doc.add_paragraph()
        skills_title.add_run('SKILLS').bold = True
        skills_title.runs[0].font.size = Pt(12)
        self._add_horizontal_line(skills_title)
        for group in self.data.get('skills', []):
            para = doc.add_paragraph()
            para.add_run(f"{group.get('name')}: {', '.join(group.get('items', []))}")

        # Projects
        proj_title = doc.add_paragraph()
        proj_title.add_run('PROJECTS').bold = True
        proj_title.runs[0].font.size = Pt(12)
        self._add_horizontal_line(proj_title)
        for proj in self.data.get('projects', []):
            p = doc.add_paragraph()
            run = p.add_run(proj.get('name', ''))
            run.bold = True
            for bullet in proj.get('bullets', []):
                self._add_bulleted_item(doc, bullet)

        doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python create_resume_docx.py resume.json [output.docx]")
        sys.exit(1)
    json_path = Path(sys.argv[1])
    output_file = Path(sys.argv[2]) if len(sys.argv) >= 3 else Path("Resume.docx")
    if not json_path.exists():
        sys.exit(1)
    with json_path.open() as f:
        data = json.load(f)
    generator = ResumeJsonToDocx(data)
    generator.generate(str(output_file))
